#!/usr/bin/python 
# coding:utf-8
"""
File: trans-docx.py
Fuction: Translate English documents to Japanese.
Version: 0.1 (As of Sep.2019)
"""
import sys
import shutil
import boto3
import docx

SOURCE_LANG = "en"
TARGET_LANG = "ja"
translate = boto3.client(service_name='translate')
# ----------------------------------------------------
# cp_file(arg1, arg2)
#     arg1: Original English Excel file name.
#     arg2: Translate target Japanese Excel file name.
#     return: none.
# ----------------------------------------------------
def cp_file(fn_in, fn_out):
    shutil.copy(fn_in, fn_out)
    return()
# ----------------------------------------------------
# trans_docx(arg)
#    arg: input sentence for translate. Maybe English
#    return: translated sentence. Maybe Japanese
# ----------------------------------------------------
def trans_docx(phrase):
    result = translate.translate_text(Text=str(phrase), SourceLanguageCode=SOURCE_LANG, TargetLanguageCode=TARGET_LANG)
    return(result.get('TranslatedText'))
# ----------------------------------------------------
# proc_word(arg)
#    arg: File name for writing traslated sentence.
# ----------------------------------------------------
def proc_docx(fn):
    doc = docx.Document(fn)
    #-------------------------------
    # Translate & Replace Paragraphs
    #-------------------------------
    #-debug- txt = []
    for par in doc.paragraphs:
        if len(par.text) == 0:
            continue
        #-debug- txt.append(par.text)
        jptxt = trans_docx(par.text)
        par.text = par.text.replace(par.text, jptxt)
        print('trans-docx : Translating text...')
    #-------------------------------
    # Translate & Replace Table text
    #-------------------------------
    paragraphs = (paragraph
                  for table in doc.tables
                  for row in table.rows
                  for cell in row.cells
                  for paragraph in cell.paragraphs)
    for paragraph in paragraphs:
        if len(paragraph.text) == 0:
            continue
        jptbl = trans_docx(paragraph.text)
        paragraph.text = paragraph.text.replace(paragraph.text, jptbl)
        print('trans-docx : Translating Table text...')
    
    doc.save(fn)
    # print(txt)
    return()
# ----------------------------------------------------
# main() 
#    arg1: Original Word file name.
#    arg2: Target Word file name. File created if there not exist. Overwrite if there is exist.
# ----------------------------------------------------
def main():
    args = sys.argv
    num = len(args)
    if num != 3:
        print('trans: ERROR: Usage $./trans_word.py <Input Word name> <Output Word name>')
        sys.exit('trans_word: ERROR.')
    rc = cp_file(args[1], args[2])  # cp & open output file.
    proc_docx(args[2])
    sys.exit('trans_docx: END Successfully.')
    
# ----------------------------------------------------
if __name__ == '__main__':
    main()
# ----------------------------------------------------
# End of File
# ----------------------------------------------------
